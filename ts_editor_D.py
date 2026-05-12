"""
TermSheetTransformer - Version D ultime

Transforme des documents Word (term sheets) avec :
- Version mark-up : modifications visibles (jaune + barré)
- Version finale : document modifié sans marquage
- Version PDF : conversion du document final

Règles de marquage :
- Remplacement (Mot1 → Mot2) : Mot1 barré + Mot2 surligné jaune
- Ajout (nouvelle section) : surligné jaune
- Modification (description) : surligné jaune
- Suppression : texte barré (conservé pour trace)
"""

import subprocess
import shutil
import sys
import tempfile
import re
import unicodedata
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


# ---op--------------------------------------------------------------------------
# Conversion .doc → .docx
# -----------------------------------------------------------------------------

def _strip_all_highlights(doc: Document, keep_red: bool = True) -> None:
    """Supprime tout surlignage du document. Si keep_red=True, conserve le surlignage rouge."""
    def clear_run(r):
        if keep_red and r.font.highlight_color == WD_COLOR_INDEX.RED:
            return
        r.font.highlight_color = None

    for p in doc.paragraphs:
        for r in p.runs:
            clear_run(r)
    for table in doc.tables:
        _strip_highlights_in_table(table, clear_run)


def _strip_highlights_in_table(table, clear_run) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    clear_run(r)
            for nested in cell.tables:
                _strip_highlights_in_table(nested, clear_run)


def _convert_doc_to_docx(doc_path: Path) -> Tuple[Path, Path]:
    """Convertit .doc en .docx. Essaie Word (Windows) puis LibreOffice."""
    tmpdir = Path(tempfile.mkdtemp(prefix="ts_transformer_"))
    docx_path = tmpdir / (doc_path.stem + ".docx")

    # 1. Windows : Microsoft Word via COM
    if sys.platform == "win32":
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(doc_path.resolve()))
            doc.SaveAs2(str(docx_path), FileFormat=16)  # 16 = wdFormatDocumentDefault (.docx)
            doc.Close()
            word.Quit()
            if docx_path.exists():
                return docx_path, tmpdir
        except Exception:
            pass

    # 2. LibreOffice (soffice / libreoffice)
    cmd = None
    for exe in ("soffice", "libreoffice"):
        found = shutil.which(exe)
        if found:
            cmd = [found, "--headless", "--convert-to", "docx", "--outdir", str(tmpdir), str(doc_path)]
            break
    if cmd is None and Path("/Applications/LibreOffice.app/Contents/MacOS/soffice").exists():
        cmd = ["/Applications/LibreOffice.app/Contents/MacOS/soffice", "--headless",
               "--convert-to", "docx", "--outdir", str(tmpdir), str(doc_path)]
    if cmd is not None:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode == 0 and docx_path.exists():
            return docx_path, tmpdir

    shutil.rmtree(tmpdir, ignore_errors=True)
    raise RuntimeError(
        "Conversion .doc impossible : ni Word (Windows) ni LibreOffice détecté. "
        "Convertissez manuellement en .docx (Word : Fichier > Enregistrer sous > Word Document)."
    )


# -----------------------------------------------------------------------------
# Enregistrements des modifications
# -----------------------------------------------------------------------------

@dataclass
class ReplaceOp:
    old: str
    new: str
    occurrence: int = 1


@dataclass
class AddSectionOp:
    after_section: str
    title: str
    description: str
    occurrence: int = 1


@dataclass
class UpdateDescriptionOp:
    section_title: str
    new_description: str
    occurrence: int = 1


@dataclass
class SetSectionOp:
    """Ajoute ou met à jour une section (fusion add_section_after + update_section_description)."""
    title: str
    description: str
    after_section: Optional[str] = None  # Si None, ajoute à la fin du tableau
    red_highlight_in_final: bool = False
    occurrence: int = 1


@dataclass
class DeleteSectionOp:
    section_title: str
    occurrence: int = 1


@dataclass
class SetSectionOrderOp:
    """Définit l'ordre des sections dans le tableau."""
    section_order: List[str]


@dataclass
class AddLogoOp:
    logo_path: str
    width_inches: float = 1.0
    all_sections: bool = True


@dataclass
class AddContentOp:
    """Ajoute un titre ou une description (paragraphe hors tableau)."""
    text: str
    after_paragraph: Optional[str] = None  # Texte du paragraphe après lequel insérer (None = fin)
    red_highlight_in_final: bool = False


@dataclass
class RemoveParagraphOp:
    """Supprime un paragraphe contenant le texte donné."""
    text_contains: str
    occurrence: int = 1


@dataclass
class UpdateParagraphOp:
    """Modifie le contenu d'un paragraphe."""
    text_contains: str
    new_text: str
    occurrence: int = 1


@dataclass
class SetDisclaimerSectionOp:
    """Ajoute ou met à jour une section dans la partie Disclaimer (après les tableaux)."""
    title: str
    content: str  # Peut contenir \n pour plusieurs paragraphes
    after_title: Optional[str] = None  # Titre après lequel insérer (None = fin)
    red_highlight_in_final: bool = False
    occurrence: int = 1


@dataclass
class RemoveDisclaimerSectionOp:
    """Supprime une section de disclaimer complète (titre + contenu)."""
    title: str
    occurrence: int = 1


@dataclass
class UpdateDisclaimerContentOp:
    """Met à jour le contenu d'une section de disclaimer existante."""
    title: str
    new_content: str  # Peut contenir \n pour plusieurs paragraphes
    red_highlight_in_final: bool = False
    occurrence: int = 1


@dataclass
class AddDisclaimerContentOp:
    """Ajoute du contenu à la fin d'une section de disclaimer existante."""
    title: str
    additional_content: str  # Peut contenir \n pour plusieurs paragraphes
    red_highlight_in_final: bool = False
    occurrence: int = 1


# -----------------------------------------------------------------------------
# Éditeur de base
# -----------------------------------------------------------------------------

class _TermSheetEditor:
    """
    Éditeur interne pour term sheets.
    markup_mode=True : applique surlignage jaune et barré sur les modifications.
    """

    def __init__(self, doc: Document, markup_mode: bool = False):
        self.doc = doc
        self.markup_mode = markup_mode

    def replace_text(self, old: str, new: str):
        for p in self.doc.paragraphs:
            self._replace_in_paragraph_runs(p, old, new)
        for t in self.doc.tables:
            self._replace_in_table(t, old, new)
        return self

    def _replace_in_table(self, table, old: str, new: str):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    self._replace_in_paragraph_runs(p, old, new)
                for nested in cell.tables:
                    self._replace_in_table(nested, old, new)

    def _replace_in_paragraph_runs(self, paragraph, old: str, new: str):
        runs = list(paragraph.runs)
        if not runs:
            return

        full_text = "".join(r.text for r in runs)
        if old not in full_text:
            return

        char_to_run = []
        for i, r in enumerate(runs):
            char_to_run.extend([i] * len(r.text))

        segments = []
        i = 0
        L = len(full_text)
        old_len = len(old)

        while i < L:
            if full_text.startswith(old, i):
                start_run_idx = char_to_run[i] if i < len(char_to_run) else 0
                if self.markup_mode:
                    segments.append(("strike", old, start_run_idx))
                    segments.append(("highlight", new, start_run_idx))
                else:
                    segments.append((new, start_run_idx))
                i += old_len
            else:
                run_idx = char_to_run[i] if i < len(char_to_run) else 0
                segments.append((full_text[i], run_idx))
                i += 1

        merged = []
        for seg in segments:
            if seg[0] in ("strike", "highlight"):
                merged.append(seg)
            else:
                text, ridx = seg
                if merged and merged[-1][0] not in ("strike", "highlight") and merged[-1][1] == ridx:
                    merged[-1] = (merged[-1][0] + text, ridx)
                else:
                    merged.append(seg)

        for r in runs:
            r._element.getparent().remove(r._element)

        for seg in merged:
            if seg[0] == "strike":
                _, text, ridx = seg
                run = paragraph.add_run(text)
                self._copy_run_format(runs[ridx], run)
                run.font.strike = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif seg[0] == "highlight":
                _, text, ridx = seg
                run = paragraph.add_run(text)
                self._copy_run_format(runs[ridx], run)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                text, ridx = seg
                if text == "":
                    continue
                run = paragraph.add_run(text)
                self._copy_run_format(runs[ridx], run)

    def _copy_run_format(self, src_run, dst_run):
        """Copie le format via XML (rPr) pour préserver tous les styles, y compris hérités."""
        src_rpr = src_run._element.find(qn("w:rPr"))
        if src_rpr is not None:
            dst_rpr = dst_run._element.find(qn("w:rPr"))
            if dst_rpr is not None:
                dst_run._element.remove(dst_rpr)
            dst_run._element.insert(0, deepcopy(src_rpr))
        else:
            # Fallback : copie propriété par propriété
            dst_run.bold = src_run.bold
            dst_run.italic = src_run.italic
            dst_run.underline = src_run.underline
            dst_run.font.name = src_run.font.name
            dst_run.font.size = src_run.font.size
            if src_run.font.color is not None and src_run.font.color.rgb is not None:
                dst_run.font.color.rgb = src_run.font.color.rgb
        if not self.markup_mode:
            dst_run.font.highlight_color = None

    def _section_block_rows(self, table, title_row):
        """
        Lignes consécutives d'une même section lorsque la colonne titre (A) n'est pas fusionnée :
        première ligne = titre en A ; lignes suivantes (à 2 tc) avec A vide = suite de la description.

        On stoppe sur :
        - une ligne fusionnée (1 tc) qui marque toujours une rupture de section (en-tête),
        - une ligne dont la colonne A contient du texte (nouveau titre).
        """
        rows_list = table.rows
        start_idx = None
        for i, r in enumerate(rows_list):
            if r._tr is title_row._tr:
                start_idx = i
                break
        if start_idx is None:
            return [title_row]
        block = []
        for j in range(start_idx, len(rows_list)):
            r = rows_list[j]
            if j == start_idx:
                block.append(r)
                continue
            if self._real_tc_count(r) != 2:
                break
            first = r.cells[0].text if len(r.cells) >= 1 else ""
            if self._normalize(first):
                break
            block.append(r)
        return block

    def _aggregate_description_from_block(self, block_rows) -> str:
        parts = []
        for r in block_rows:
            if self._real_tc_count(r) == 2:
                parts.append(r.cells[1].text)
        return "\n".join(parts).strip()

    def _apply_description_to_block(
        self, block_rows, description: str, highlight: bool, red_highlight: bool
    ):
        """Écrit toute la description dans la première cellule B du bloc et vide les B suivantes.

        N'écrit dans `cells[1]` que si la ligne a 2 `<w:tc>` réels (titre et
        description physiquement distincts). Pour une ligne fusionnée (1 tc,
        gridSpan=2), `cells[0]` et `cells[1]` pointent vers la même cellule
        physique : écrire dans `cells[1]` écraserait le titre. On saute donc
        ces lignes.
        """
        if not block_rows:
            return
        first_desc = True
        for r in block_rows:
            if self._real_tc_count(r) != 2:
                continue
            if first_desc:
                self._set_cell_text(
                    r.cells[1],
                    description,
                    highlight=highlight,
                    red_highlight=red_highlight,
                )
                first_desc = False
            else:
                self._set_cell_text(r.cells[1], "", highlight=False, red_highlight=False)

    def update_section_description(self, section_title: str, new_description: str, occurrence: int = 1):
        row, table = self._find_section_row(section_title, occurrence)
        if row is None:
            return self
        if self._real_tc_count(row) == 2:
            block = self._section_block_rows(table, row)
            self._apply_description_to_block(
                block, new_description, highlight=self.markup_mode, red_highlight=False
            )
        else:
            title_part, sep = self._extract_1tc_title_and_sep(row.cells[0].text)
            self._set_cell_text(
                row.cells[0],
                f"{title_part}{sep}{new_description}",
                highlight=self.markup_mode,
            )
        return self

    def delete_section(self, section_title: str, occurrence: int = 1):
        row, table = self._find_section_row(section_title, occurrence)
        if row is None:
            return self
        if self._real_tc_count(row) == 2:
            block = self._section_block_rows(table, row)
        else:
            block = [row]
        if self.markup_mode:
            for r in block:
                self._strike_row(r)
        else:
            for r in reversed(block):
                table._tbl.remove(r._tr)
        return self

    def _strike_row(self, row):
        """Barré tout le contenu d'une ligne (mode mark-up pour suppression)."""
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.strike = True
                    r.font.highlight_color = WD_COLOR_INDEX.YELLOW

    def get_section_description(self, section_title: str, occurrence: int = 1) -> Optional[str]:
        """
        Lit la description d'une section et la retourne sous forme de string.
        - Retourne None si la section est introuvable.
        - Pour les lignes à 2 colonnes : concatène les cellules B des lignes consécutives dont la
          colonne A est vide après la ligne titre (cas sans fusion verticale).
        - Pour les lignes à 1 colonne (format "Titre: description") : retourne la partie après le premier ':'.
        """
        row, table = self._find_section_row(section_title, occurrence)
        if row is None:
            return None
        if self._real_tc_count(row) == 2:
            block = self._section_block_rows(table, row)
            return self._aggregate_description_from_block(block)
        text = row.cells[0].text if len(row.cells) >= 1 else ""
        if ":" in text:
            return text.split(":", 1)[1].strip()
        return None

    def insert_section_after(self, after_section: str, new_title: str, new_description: str, occurrence: int = 1):
        ref_row, table = self._find_section_row(after_section, occurrence)
        if ref_row is None:
            raise ValueError(f"Section de référence introuvable : {after_section}")
        anchor = (
            self._section_block_rows(table, ref_row)[-1]
            if self._real_tc_count(ref_row) == 2
            else ref_row
        )
        self._create_minimal_row_after(
            table, ref_row, new_title, new_description, insert_after_row=anchor
        )
        return self

    def set_section(
        self, 
        title: str, 
        description: str, 
        after_section: Optional[str] = None,
        red_highlight: bool = False,
        occurrence: int = 1
    ):
        """
        Ajoute ou met à jour une section.
        - Si la section existe : met à jour la description (tout le texte est regroupé dans la
          première cellule B du bloc ; les lignes de continuation avec colonne A vide sont vidées en B)
        - Si la section n'existe pas : l'ajoute après after_section (ou à la fin si None).
          Si after_section désigne une section multi-lignes sans fusion, l'insertion se fait après
          la dernière ligne de ce bloc.
        """
        row, table = self._find_section_row(title, occurrence)

        if row is not None:
            tc_count = self._real_tc_count(row)
            if tc_count == 2:
                block = self._section_block_rows(table, row)
                self._apply_description_to_block(
                    block,
                    description,
                    highlight=self.markup_mode,
                    red_highlight=red_highlight and not self.markup_mode,
                )
            else:
                # Ligne fusionnée (1 tc) : on remplace l'ancienne description en
                # conservant le titre original. On NE touche jamais cells[1] ici car
                # ce serait la même cellule physique que cells[0].
                title_part, sep = self._extract_1tc_title_and_sep(row.cells[0].text)
                self._set_cell_text(
                    row.cells[0],
                    f"{title_part}{sep}{description}",
                    highlight=self.markup_mode,
                    red_highlight=red_highlight and not self.markup_mode,
                )
        else:
            # Section inexistante : on l'ajoute toujours dans le tableau principal
            # 2-col, jamais dans un tableau externe.
            if not self.doc.tables:
                raise ValueError("Aucun tableau trouvé dans le document")
            table = self._find_main_table()
            if table is None or not table.rows:
                raise ValueError("Aucun tableau 2-colonnes trouvé pour l'insertion")
            format_row = self._find_reference_row_for_format(table)
            if format_row is None:
                raise ValueError("Impossible de trouver une ligne de référence valide dans le tableau")

            if after_section:
                ref_row, ref_table = self._find_section_row(after_section, occurrence)
                if ref_row is None:
                    raise ValueError(f"Section de référence introuvable : {after_section}")
                # On exige que la section de référence soit dans le tableau principal,
                # sinon on insère à la fin du tableau principal.
                if ref_table is not table:
                    anchor = table.rows[-1]
                else:
                    anchor = (
                        self._section_block_rows(table, ref_row)[-1]
                        if self._real_tc_count(ref_row) == 2
                        else ref_row
                    )
            else:
                anchor = table.rows[-1]

            self._create_minimal_row_after(
                table,
                format_row,
                title,
                description,
                red_highlight,
                insert_after_row=anchor,
            )
        return self

    def _real_tc_count(self, row) -> int:
        """Nombre RÉEL de cellules physiques `<w:tc>` dans la ligne XML.

        Différent de `len(row.cells)` qui déplie les `gridSpan` et renvoie
        autant d'entrées que de colonnes logiques (potentiellement plusieurs
        entrées pointant vers la même cellule physique).
        """
        return len(row._tr.findall(qn("w:tc")))

    def _row_title_matches(self, row, target: str) -> bool:
        """Indique si la ligne porte le titre `target`.

        - 2 tc : `cells[0].text` doit valoir exactement `target`.
        - 1 tc : `cells[0].text` doit valoir `target`, ou commencer par
          `target + ":"` / `target + " :"` (format "Titre : description" dans
          la même cellule pour les sections fusionnées).
        """
        if not row.cells:
            return False
        tc = self._real_tc_count(row)
        cell_text = self._normalize(row.cells[0].text)
        if tc == 2:
            return cell_text == target
        if tc == 1:
            if cell_text == target:
                return True
            return (
                cell_text.startswith(target + ":")
                or cell_text.startswith(target + " :")
            )
        return False

    def _extract_1tc_title_and_sep(self, cell_text: str) -> Tuple[str, str]:
        """Pour une cellule 1-tc, extrait (titre, séparateur).

        Renvoie le titre nu (avant `:`) et le séparateur utilisé (`" : "` ou
        `": "`), pour préserver le style original lors d'une mise à jour.
        """
        normalized = self._normalize(cell_text)
        if " :" in normalized:
            head, _ = normalized.split(":", 1)
            head = head.rstrip()
            return head, " : "
        if ":" in normalized:
            head, _ = normalized.split(":", 1)
            return head.rstrip(), ": "
        return normalized, ": "

    def _grid_col_count(self, table) -> int:
        """Nombre de colonnes canoniques du tableau d'après `<w:tblGrid>`."""
        tbl = table._tbl
        grid = tbl.find(qn("w:tblGrid"))
        if grid is None:
            return 0
        return len(grid.findall(qn("w:gridCol")))

    def _is_two_col_table(self, table) -> bool:
        """
        Détecte si un tableau est un tableau "titre/description" à 2 colonnes.

        Critère strict combinant plusieurs garde-fous :
        - `<w:tblGrid>` doit avoir au plus 2 colonnes canoniques (un tableau
          5-col à `tblGrid` 5 est immédiatement exclu, même si ses lignes ont
          été fusionnées en 2 tc visibles).
        - Aucune ligne ne doit avoir plus de 2 `<w:tc>`.
        - Au moins une ligne doit avoir exactement 2 `<w:tc>` (lignes
          titre/description).
        """
        if not table.rows:
            return False
        grid_cols = self._grid_col_count(table)
        if grid_cols > 2:
            return False
        has_two_tc_row = False
        for row in table.rows:
            tc_count = self._real_tc_count(row)
            if tc_count > 2:
                return False
            if tc_count == 2:
                has_two_tc_row = True
        return has_two_tc_row

    def _find_main_table(self):
        """
        Retourne le tableau principal du document :
        le tableau "2-colonnes" titre/description ayant le plus grand nombre
        de lignes à 2 `<w:tc>` réels (vraies lignes titre/description).

        Les tableaux dont `<w:tblGrid>` a >2 colonnes ou dont une ligne a >2 tc
        sont exclus. Repli sur doc.tables[0] si aucun candidat n'est trouvé.
        """
        best = None
        best_score = -1
        for table in self.doc.tables:
            if not self._is_two_col_table(table):
                continue
            score = sum(1 for r in table.rows if self._real_tc_count(r) == 2)
            if score > best_score:
                best = table
                best_score = score
        return best if best is not None else (self.doc.tables[0] if self.doc.tables else None)

    def set_section_order(self, section_order: List[str]):
        """
        Réorganise les sections du document selon l'ordre demandé, en gérant
        les tableaux non-2-col (ex: tableau 5-col de description) attachés à
        une ligne de titre.

        Principe :
        - On considère TOUS les tableaux 2-col du body en ordre.
        - Chaque ligne d'un tableau 2-col est une "section item".
        - Les éléments du body qui ne sont PAS des tableaux 2-col (paragraphe
          ou tableau non-2-col) et qui sont entre le premier et le dernier
          tableau 2-col sont "attachés" à la section qui les précède.
        - Un "bloc" = ligne de titre + lignes de continuation (colonne A vide
          juste après). Tout le bloc se déplace comme une unité, ainsi que ses
          externals attachés.
        - Les blocs dont le titre est dans `section_order` sont triés ; les
          autres restent à leur slot d'origine.
        - Les éléments AVANT le premier tableau 2-col (pre-header) et ceux
          APRÈS le dernier (post-footer) restent intouchés.
        - Si une section déplacée a un tableau 5-col attaché, il la suit dans
          le body : le tableau 2-col cible est scindé pour insérer le 5-col
          juste après la ligne propriétaire.
        """
        if not self.doc.tables:
            return self

        body = self.doc.element.body
        body_children = list(body)

        # 1. Repérer tous les tableaux 2-col top-level
        two_col_tbl_elements = []
        two_col_tbl_docs = []
        for child in body_children:
            if child.tag != qn("w:tbl"):
                continue
            tbl_doc = next(
                (t for t in self.doc.tables if t._tbl is child), None
            )
            if tbl_doc is not None and self._is_two_col_table(tbl_doc):
                two_col_tbl_elements.append(child)
                two_col_tbl_docs.append(tbl_doc)

        if not two_col_tbl_docs:
            return self

        first_idx = body_children.index(two_col_tbl_elements[0])
        last_2col_idx = body_children.index(two_col_tbl_elements[-1])

        # On étend la zone managée si des tableaux non-2-col suivent le
        # dernier tableau 2-col : ils sont attachés à la dernière ligne et
        # doivent suivre celle-ci lors d'un déplacement.
        last_managed_idx = last_2col_idx
        for i in range(last_2col_idx + 1, len(body_children)):
            if body_children[i].tag == qn("w:tbl"):
                last_managed_idx = i

        pre_header_elements = body_children[:first_idx]
        managed_elements = body_children[first_idx:last_managed_idx + 1]
        post_footer_elements = body_children[last_managed_idx + 1:]

        # 2. Construire la séquence à plat de "row" / "external" dans la zone managée
        items = []
        two_col_set = set(id(e) for e in two_col_tbl_elements)
        for child in managed_elements:
            if id(child) in two_col_set:
                tbl_doc = two_col_tbl_docs[two_col_tbl_elements.index(child)]
                for row in tbl_doc.rows:
                    items.append({"type": "row", "row": row, "tr": row._tr})
            else:
                items.append({"type": "external", "element": child})

        # 3. Sections : chaque ligne + ses externals juste après
        sections = []
        for it in items:
            if it["type"] == "row":
                sections.append({"row_item": it, "attached": []})
            else:
                if sections:
                    sections[-1]["attached"].append(it)

        if not sections:
            return self

        # 4. Blocs : titre + lignes de continuation (A vide)
        blocks = []
        current_block = None
        for sec in sections:
            row = sec["row_item"]["row"]
            tc = self._real_tc_count(row)
            cell_text = self._normalize(row.cells[0].text if row.cells else "")
            if tc == 1:
                clean_title, _ = self._extract_1tc_title_and_sep(row.cells[0].text)
            else:
                clean_title = cell_text

            if cell_text:
                current_block = {"sections": [sec], "clean_title": clean_title}
                blocks.append(current_block)
            else:
                if current_block is None:
                    current_block = {"sections": [sec], "clean_title": ""}
                    blocks.append(current_block)
                else:
                    current_block["sections"].append(sec)

            # Si cette section a des externals attachés (ex: 5-col après une
            # ligne titre d'un tableau 2-col), on clôt le bloc courant. La
            # prochaine section démarrera un nouveau bloc même si sa colonne A
            # est vide (cas : 2-col table suivante commençant par une ligne de
            # continuation orpheline).
            if sec["attached"]:
                current_block = None

        if not blocks:
            return self

        # 5. Marquer named vs anchor
        order_map = {
            self._normalize(t): rank for rank, t in enumerate(section_order)
        }
        for idx, b in enumerate(blocks):
            b["orig_idx"] = idx
            b["is_named"] = b["clean_title"] in order_map
            b["rank"] = order_map.get(b["clean_title"], len(order_map))

        if not any(b["is_named"] for b in blocks):
            return self

        # 6. Trier les blocs nommés dans leurs slots d'origine
        named_slot_indices = [i for i, b in enumerate(blocks) if b["is_named"]]
        named_blocks_sorted = sorted(
            [b for b in blocks if b["is_named"]],
            key=lambda b: (b["rank"], b["orig_idx"]),
        )
        new_blocks = list(blocks)
        for slot_i, sb in zip(named_slot_indices, named_blocks_sorted):
            new_blocks[slot_i] = sb

        # 7. Reconstruction : on crée de nouveaux tableaux 2-col en clonant
        #    le template (le premier 2-col du document) et on y replace les
        #    <w:tr> originaux. Les externals attachés cassent le tableau en
        #    cours et démarrent un nouveau tableau après.
        template_tbl = two_col_tbl_elements[0]

        def make_new_2col_tbl():
            new_tbl = deepcopy(template_tbl)
            for tr in list(new_tbl.findall(qn("w:tr"))):
                new_tbl.remove(tr)
            return new_tbl

        new_body_elements = []
        current_new_tbl = None
        for b in new_blocks:
            for sec in b["sections"]:
                tr = sec["row_item"]["tr"]
                tr_parent = tr.getparent()
                if tr_parent is not None:
                    tr_parent.remove(tr)
                if current_new_tbl is None:
                    current_new_tbl = make_new_2col_tbl()
                    new_body_elements.append(current_new_tbl)
                current_new_tbl.append(tr)

                for ext_item in sec["attached"]:
                    ext = ext_item["element"]
                    ext_parent = ext.getparent()
                    if ext_parent is not None:
                        ext_parent.remove(ext)
                    new_body_elements.append(ext)
                    current_new_tbl = None

        # 8. Retirer les anciens tableaux 2-col (devenus vides après le move)
        for tbl_elem in two_col_tbl_elements:
            if tbl_elem.find(qn("w:tr")) is None:
                tp = tbl_elem.getparent()
                if tp is not None:
                    tp.remove(tbl_elem)

        # 9. Insérer le nouveau contenu entre pre_header et post_footer
        if pre_header_elements:
            cursor = pre_header_elements[-1]
            for elem in new_body_elements:
                cursor.addnext(elem)
                cursor = elem
        elif post_footer_elements:
            first_post = post_footer_elements[0]
            for elem in reversed(new_body_elements):
                first_post.addprevious(elem)
        else:
            for elem in new_body_elements:
                body.append(elem)

        return self

    def _set_cell_text(self, cell, text: str, highlight: bool = False, red_highlight: bool = False):
        """
        Définit le texte d'une cellule en gérant les retours à la ligne (\n).
        """
        # Gérer les retours à la ligne
        lines = text.split('\n')
        
        if not cell.paragraphs:
            cell.add_paragraph("")
        
        # Supprimer tous les paragraphes sauf le premier
        for extra_p in cell.paragraphs[1:]:
            extra_p._element.getparent().remove(extra_p._element)
        
        # Premier paragraphe avec la première ligne
        p = cell.paragraphs[0]
        ref_run = None
        if p.runs:
            ref_run = p.runs[0]
            ref_run.text = lines[0] if lines else ""
            if highlight:
                ref_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif red_highlight:
                ref_run.font.highlight_color = WD_COLOR_INDEX.RED
            elif not self.markup_mode:
                ref_run.font.highlight_color = None
            # Supprimer les autres runs
            for r in p.runs[1:]:
                r.text = ""
        else:
            r = p.add_run(lines[0] if lines else "")
            if highlight:
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif red_highlight:
                r.font.highlight_color = WD_COLOR_INDEX.RED
            elif not self.markup_mode:
                r.font.highlight_color = None
            ref_run = r
        
        # Ajouter les lignes suivantes comme nouveaux paragraphes
        for line in lines[1:]:
            new_p = cell.add_paragraph("")
            new_run = new_p.add_run(line)

            # Copier le format du premier run pour préserver police/taille/etc.
            if ref_run is not None:
                self._copy_run_format(ref_run, new_run)

            if new_p.runs:
                if highlight:
                    new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                elif red_highlight:
                    new_run.font.highlight_color = WD_COLOR_INDEX.RED
                elif not self.markup_mode:
                    new_run.font.highlight_color = None

    def _normalize(self, s: str) -> str:
        return " ".join((s or "").replace("\n", " ").split()).strip()

    def _normalize_loose(self, s: str) -> str:
        """
        Normalisation souple pour la recherche de paragraphes :
        - case-insensitive
        - ignore ponctuation et multiples espaces
        """
        normalized = self._normalize(s).casefold()
        umlaut_map = {
            "ä": "ae", "ö": "oe", "ü": "ue", "ß": "ss",
            "à": "a", "á": "a", "â": "a",
            "è": "e", "é": "e", "ê": "e", "ë": "e",
            "ì": "i", "í": "i", "î": "i", "ï": "i",
            "ò": "o", "ó": "o", "ô": "o",
            "ù": "u", "ú": "u", "û": "u",
            "ç": "c",
        }
        normalized = "".join(umlaut_map.get(c, c) for c in normalized)
        normalized = "".join(
            c for c in unicodedata.normalize("NFKD", normalized)
            if not unicodedata.combining(c)
        )
        normalized = re.sub(r"[^\w\s]", " ", normalized)
        return " ".join(normalized.split()).strip()

    def _contains_normalized(self, haystack: str, needle: str) -> bool:
        """
        Vérifie une inclusion textuelle robuste :
        - exact brut
        - brut case-insensitive
        - inclusion après normalisation souple
        """
        if not needle:
            return False
        if needle in haystack:
            return True
        if needle.casefold() in haystack.casefold():
            return True
        normalized_needle = self._normalize_loose(needle)
        normalized_haystack = self._normalize_loose(haystack)
        if normalized_needle in normalized_haystack:
            return True

        # Variante plus permissive pour allemand translittéré (ae/oe/ue -> a/o/u).
        plain_needle = normalized_needle.replace("ae", "a").replace("oe", "o").replace("ue", "u")
        plain_haystack = normalized_haystack.replace("ae", "a").replace("oe", "o").replace("ue", "u")
        return plain_needle in plain_haystack

    def _paragraph_matches_query(self, paragraph, query: str) -> bool:
        """Retourne True si le paragraphe correspond à la requête de façon tolérante."""
        return self._contains_normalized(paragraph.text or "", query or "")

    def _title_key(self, s: str) -> str:
        """Clé de comparaison des titres de disclaimer (tolère ':' final et casse)."""
        return self._normalize((s or "").rstrip(":")).casefold()

    def _xml_onoff_true(self, elem, tag) -> bool:
        """
        Décode un booléen WordprocessingML (ex: w:b, w:i, w:u) en tenant compte de w:val.
        """
        if elem is None:
            return False
        node = elem.find(tag)
        if node is None:
            return False
        val = node.get(qn("w:val"))
        return val not in ("0", "false", "off")

    def _style_chain_bold(self, style):
        """
        Retourne True/False si le gras est explicitement défini dans la chaîne
        de styles (style -> base_style), sinon None.
        """
        current = style
        while current is not None:
            try:
                bold = current.font.bold
            except AttributeError:
                bold = None
            if bold is not None:
                return bool(bold)
            current = getattr(current, "base_style", None)
        return None

    def _run_is_effectively_bold(self, run, para_rpr, para_style=None) -> bool:
        """
        Détermine si un run est effectivement gras:
        - propriété directe du run
        - XML run/paragraphe
        - héritage des styles (run style / paragraphe style)
        """
        if run.bold is True:
            return True
        if run.bold is False:
            return False

        run_rpr = run._element.find(qn("w:rPr"))
        if self._xml_onoff_true(run_rpr, qn("w:b")) or self._xml_onoff_true(para_rpr, qn("w:b")):
            return True

        run_style_bold = self._style_chain_bold(getattr(run, "style", None))
        if run_style_bold is not None:
            return run_style_bold

        para_style_bold = self._style_chain_bold(para_style)
        if para_style_bold is not None:
            return para_style_bold

        return False

    def _find_reference_row_for_format(self, table):
        """
        Trouve une ligne de référence avec le bon format (exactement 2 cellules
        physiques) pour cloner. Les lignes fusionnées (1 tc, gridSpan=2) sont
        exclues.
        Priorité : "Currency" > "Trade Date" > première ligne 2-tc trouvée.
        """
        reference_titles = ["currency", "trade date"]

        def norm(s: str) -> str:
            return self._normalize(s).casefold()

        for title_to_search in reference_titles:
            for row in table.rows:
                if self._real_tc_count(row) == 2 and norm(row.cells[0].text) == title_to_search:
                    return row

        for row in table.rows:
            if self._real_tc_count(row) == 2:
                return row

        return None

    def _find_section_row(self, section_title: str, occurrence: int = 1):
        """
        Cherche une ligne de section dans les tableaux 2-colonnes uniquement.

        Les tableaux à >2 colonnes (descriptions externes) sont entièrement ignorés
        pour éviter qu'un titre de section corresponde par erreur au contenu d'une
        cellule de tableau 5-colonnes.
        """
        target = self._normalize(section_title)
        count = 0
        for table in self.doc.tables:
            if not self._is_two_col_table(table):
                continue
            row = self._find_section_row_in_table(table, target, occurrence, count)
            if row is not None:
                return row, table
            count += self._count_occurrences_in_table(table, target)
        return None, None

    def _count_occurrences_in_table(self, table, target: str) -> int:
        """Compte les occurrences du titre dans un tableau 2-col.

        Accepte les lignes 2-tc (titre dans `cells[0]`) ET les lignes 1-tc dont
        la cellule unique vaut le titre ou commence par `titre :` (sections
        fusionnées "Titre : description" dans une seule cellule).
        """
        if not self._is_two_col_table(table):
            return 0
        c = 0
        for row in table.rows:
            if self._row_title_matches(row, target):
                c += 1
            for cell in row.cells:
                for nested in cell.tables:
                    c += self._count_occurrences_in_table(nested, target)
        return c

    def _find_section_row_in_table(self, table, target: str, occurrence: int, count_so_far: int = 0):
        """Cherche la n-ième ligne au titre demandé, dans un tableau 2-col.

        Accepte les lignes 2-tc et 1-tc (cf. `_row_title_matches`).
        """
        if not self._is_two_col_table(table):
            return None
        count = count_so_far
        for row in table.rows:
            if self._row_title_matches(row, target):
                count += 1
                if count == occurrence:
                    return row
            for cell in row.cells:
                for nested in cell.tables:
                    found = self._find_section_row_in_table(nested, target, occurrence, count)
                    if found is not None:
                        return found
                    count += self._count_occurrences_in_table(nested, target)
        return None

    def _create_minimal_row_after(self, table, ref_row, new_title: str, new_description: str, red_highlight: bool = False, insert_after_row=None):
        """
        Clone la ligne (format exact), modifie uniquement le texte (w:t) sans toucher pPr/rPr.
        
        Args:
            table: Le tableau
            ref_row: Ligne de référence pour le FORMAT (sera clonée)
            new_title: Titre de la nouvelle ligne
            new_description: Description de la nouvelle ligne
            red_highlight: Surligner en rouge?
            insert_after_row: Ligne après laquelle insérer (si None, insère après ref_row)
        """
        ref_tr = ref_row._tr
        new_tr = deepcopy(ref_tr)
        
        # Insérer après la ligne spécifiée ou après ref_row par défaut
        if insert_after_row is not None:
            insert_after_row._tr.addnext(new_tr)
        else:
            ref_tr.addnext(new_tr)

        # Récupérer l'objet Row python-docx correspondant à new_tr
        new_row = None
        for row in table.rows:
            if row._tr is new_tr:
                new_row = row
                break

        if new_row is None:
            return self

        if self._real_tc_count(new_row) == 2:
            self._set_cell_text(
                new_row.cells[0],
                new_title,
                highlight=self.markup_mode,
                red_highlight=False,
            )
            self._set_cell_text(
                new_row.cells[1],
                new_description,
                highlight=self.markup_mode,
                red_highlight=red_highlight and not self.markup_mode,
            )
        else:
            # Ligne clonée à 1 tc (fusionnée) : on encode "Titre: description"
            # dans l'unique cellule physique pour éviter d'écraser le titre.
            self._set_cell_text(
                new_row.cells[0],
                f"{new_title}: {new_description}",
                highlight=self.markup_mode,
                red_highlight=red_highlight and not self.markup_mode,
            )

        # Pas besoin de retourner la nouvelle ligne
        # car elle est déjà ajoutée au tableau
        return self

    def add_content(
        self,
        text: str,
        after_paragraph: Optional[str] = None,
        highlight: bool = False,
        red_highlight: bool = False,
    ):
        """Ajoute un paragraphe (titre/description) hors tableau."""
        if after_paragraph:
            target = self._find_body_paragraph(after_paragraph, occurrence=1)
            if target is None:
                raise ValueError(f"Paragraphe introuvable contenant : {after_paragraph!r}")
            new_p = self.doc.add_paragraph(text)
            new_p._element.getparent().remove(new_p._element)
            target._element.addnext(new_p._element)
            target_ppr = target._element.find(qn("w:pPr"))
            if target_ppr is not None:
                new_ppr = new_p._element.find(qn("w:pPr"))
                if new_ppr is not None:
                    new_p._element.remove(new_ppr)
                new_p._element.insert(0, deepcopy(target_ppr))
            if target.runs and new_p.runs:
                self._copy_run_format(target.runs[-1], new_p.runs[0])
        else:
            # Créer un paragraphe temporaire pour trouver la position
            temp_p = self.doc.add_paragraph()
            last_run, last_para = self._get_last_run_and_paragraph_before(temp_p._element)
            
            # Supprimer le paragraphe temporaire
            temp_p._element.getparent().remove(temp_p._element)
            
            # Cloner le dernier paragraphe (structure complète + style) puis modifier le texte
            if last_para is not None:
                new_p_elem = deepcopy(last_para._element)
                self.doc.element.body.append(new_p_elem)
                
                # Récupérer l'objet Paragraph python-docx
                new_p = None
                for p in self.doc.paragraphs:
                    if p._element is new_p_elem:
                        new_p = p
                        break
                
                # Modifier le texte du premier run (garde le rPr avec formatage)
                # et supprimer les runs suivants
                if new_p.runs:
                    new_p.runs[0].text = text
                    for run in list(new_p.runs[1:]):
                        run._element.getparent().remove(run._element)
                else:
                    # Si pas de run, en créer un (fallback)
                    new_p.add_run(text)
            else:
                # Fallback : créer un paragraphe normal si aucun paragraphe précédent
                new_p = self.doc.add_paragraph(text)
        for run in new_p.runs:
            if highlight:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif red_highlight:
                run.font.highlight_color = WD_COLOR_INDEX.RED
            elif not self.markup_mode:
                run.font.highlight_color = None
        return self

    def _get_last_run_and_paragraph_before(self, exclude_element):
        """Retourne (dernier Run, dernier Paragraph) avant l'élément exclu (objets python-docx)."""
        body = self.doc.element.body
        last_run, last_para = None, None
        try:
            idx = list(body).index(exclude_element)
        except ValueError:
            idx = len(body)
        blocks_before = list(body)[:idx]

        for block in reversed(blocks_before):
            if block.tag == qn("w:p"):
                for p in self.doc.paragraphs:
                    if p._element is block:
                        last_para = p
                        last_run = p.runs[-1] if p.runs else None
                        return last_run, last_para
            elif block.tag == qn("w:tbl"):
                for table in self.doc.tables:
                    if table._tbl is block:
                        run, para = self._get_last_run_para_in_table(table)
                        return (run, para) if (run or para) else (last_run, last_para)
        return last_run, last_para

    def _get_last_run_para_in_table(self, table):
        """Dernier (Run, Paragraph) dans un tableau (objets python-docx)."""
        last_run, last_para = None, None
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    last_para = p
                    last_run = p.runs[-1] if p.runs else None
                for nested in cell.tables:
                    r, p = self._get_last_run_para_in_table(nested)
                    if r is not None:
                        last_run = r
                    if p is not None:
                        last_para = p
        return last_run, last_para

    def _copy_run_format_from_api(self, src_run, dst_run):
        """Copie le format via l'API python-docx (gère l'héritage des styles)."""
        try:
            dst_run.bold = src_run.bold
            dst_run.italic = src_run.italic
            dst_run.underline = src_run.underline
            if src_run.font.name:
                dst_run.font.name = src_run.font.name
            if src_run.font.size:
                dst_run.font.size = src_run.font.size
            if src_run.font.color and src_run.font.color.rgb:
                dst_run.font.color.rgb = src_run.font.color.rgb
        except (AttributeError, TypeError):
            pass

    def remove_paragraph(self, text_contains: str, occurrence: int = 1):
        """Supprime un paragraphe hors tableau contenant le texte donné.
        Mode mark-up : barre le texte au lieu de le supprimer."""
        target = self._find_body_paragraph(text_contains, occurrence)
        if target is None:
            return self
        if self.markup_mode:
            for run in target.runs:
                run.font.strike = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            target._element.getparent().remove(target._element)
        return self

    def update_paragraph(self, text_contains: str, new_text: str, occurrence: int = 1):
        """Modifie le contenu d'un paragraphe hors tableau.
        Mode mark-up : surligne le nouveau texte en jaune."""
        target = self._find_body_paragraph(text_contains, occurrence)
        if target is None:
            return self
        target.clear()
        run = target.add_run(new_text)
        if self.markup_mode:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return self

    def _find_body_paragraph(self, text_contains: str, occurrence: int = 1):
        """Trouve le n-ième paragraphe du body (hors tableaux) avec matching robuste."""
        count = 0
        for p in self.doc.paragraphs:
            if self._paragraph_matches_query(p, text_contains):
                count += 1
                if count == occurrence:
                    return p
        return None

    # -------------------------------------------------------------------------
    # Gestion des sections de Disclaimer (paragraphes après les tableaux)
    # -------------------------------------------------------------------------

    def _is_disclaimer_title(self, paragraph) -> bool:
        """
        Détecte si un paragraphe est un titre de disclaimer.
        Critère robuste :
        - texte non vide
        - aucun run souligné
        - majoritairement gras (incluant le gras hérité des styles Word)
        """
        text = self._normalize(paragraph.text)
        if not text:
            return False

        non_empty_runs = [r for r in paragraph.runs if r.text.strip()]
        if not non_empty_runs:
            return False

        # Exclure les titres soulignés (ex: "Risk Terms") qui structurent autrement le document.
        if any(r.underline for r in non_empty_runs):
            return False

        ppr = paragraph._element.find(qn("w:pPr"))
        para_rpr = ppr.find(qn("w:rPr")) if ppr is not None else None
        para_style = getattr(paragraph, "style", None)

        total_chars = 0
        bold_chars = 0
        for run in non_empty_runs:
            run_text = run.text or ""
            char_count = len(run_text)
            total_chars += char_count
            if self._run_is_effectively_bold(run, para_rpr, para_style=para_style):
                bold_chars += char_count

        if total_chars == 0:
            return False

        # Tolérance utile quand Word segmente le texte en runs hétérogènes.
        bold_ratio = bold_chars / total_chars
        if bold_ratio < 0.8:
            return False

        # Garde-fou : un titre de section est en général court.
        return len(text) <= 140

    def _get_last_table_index(self):
        """Retourne l'index du dernier tableau dans le body."""
        last_table_idx = -1
        for i, elem in enumerate(self.doc.element.body):
            if elem.tag.split('}')[-1] == 'tbl':
                last_table_idx = i
        return last_table_idx

    def _get_disclaimer_paragraphs(self):
        """Retourne tous les paragraphes après le dernier tableau."""
        last_table_idx = self._get_last_table_index()
        
        disclaimer_paras = []
        for i, elem in enumerate(self.doc.element.body):
            if i > last_table_idx and elem.tag.split('}')[-1] == 'p':
                for p in self.doc.paragraphs:
                    if p._element == elem:
                        disclaimer_paras.append(p)
                        break
        return disclaimer_paras

    def _find_disclaimer_section(self, title: str, occurrence: int = 1):
        """
        Trouve une section de disclaimer par son titre.
        Retourne: (titre_paragraph, content_paragraphs[], last_content_paragraph)
        """
        disclaimer_paras = self._get_disclaimer_paragraphs()
        target = self._title_key(title)

        count = 0
        for i, para in enumerate(disclaimer_paras):
            # On privilégie le matching textuel du titre (plus robuste que le style seul).
            if self._title_key(para.text) == target:
                count += 1
                if count == occurrence:
                    # Récupérer le contenu (paragraphes suivants jusqu'au prochain titre)
                    content = []
                    last_content = None
                    for j in range(i + 1, len(disclaimer_paras)):
                        if self._is_disclaimer_title(disclaimer_paras[j]):
                            break
                        if disclaimer_paras[j].text.strip():
                            content.append(disclaimer_paras[j])
                            last_content = disclaimer_paras[j]
                    return para, content, last_content

        return None, [], None

    def _get_last_disclaimer_paragraph(self):
        """Retourne le dernier paragraphe de la zone disclaimer."""
        disclaimer_paras = self._get_disclaimer_paragraphs()
        return disclaimer_paras[-1] if disclaimer_paras else None

    def set_disclaimer_section(
        self,
        title: str,
        content: str,
        after_title: Optional[str] = None,
        red_highlight: bool = False,
        occurrence: int = 1
    ):
        """
        Ajoute ou met à jour une section dans la partie Disclaimer.
        - Si la section existe : remplace le contenu
        - Si la section n'existe pas : l'ajoute après after_title (ou à la fin)
        
        Le content peut contenir des \n pour séparer les paragraphes.
        """
        title_para, content_paras, _ = self._find_disclaimer_section(title, occurrence)
        
        if title_para is not None:
            # Section existe : remplacer le contenu
            if self.markup_mode:
                # Barrer l'ancien contenu
                for cp in content_paras:
                    for r in cp.runs:
                        r.font.strike = True
                        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                # Supprimer l'ancien contenu
                for cp in content_paras:
                    cp._element.getparent().remove(cp._element)
            
            # Ajouter le nouveau contenu après le titre
            content_lines = content.split('\n')
            last_elem = title_para._element
            ref_run = content_paras[0].runs[0] if content_paras and content_paras[0].runs else None
            
            for line in content_lines:
                if line.strip():
                    new_p = self.doc.add_paragraph()
                    new_p._element.getparent().remove(new_p._element)
                    last_elem.addnext(new_p._element)
                    last_elem = new_p._element
                    
                    # Copier le format du premier paragraphe de contenu
                    if content_paras and content_paras[0]._element.find(qn("w:pPr")) is not None:
                        ppr = content_paras[0]._element.find(qn("w:pPr"))
                        new_ppr = new_p._element.find(qn("w:pPr"))
                        if new_ppr is not None:
                            new_p._element.remove(new_ppr)
                        new_p._element.insert(0, deepcopy(ppr))
                    
                    # Ajouter le texte
                    run = new_p.add_run(line)
                    if ref_run:
                        self._copy_run_format(ref_run, run)
                    
                    # Appliquer highlighting
                    if self.markup_mode:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    elif red_highlight:
                        run.font.highlight_color = WD_COLOR_INDEX.RED
        else:
            # Section n'existe pas : créer titre + contenu
            # Trouver une section de disclaimer existante pour copier le style
            ref_title_for_style = None
            ref_content_for_style = None
            
            disclaimer_paras = self._get_disclaimer_paragraphs()
            for para in disclaimer_paras:
                if self._is_disclaimer_title(para) and ref_title_for_style is None:
                    ref_title_for_style = para
                elif not self._is_disclaimer_title(para) and para.text.strip() and ref_content_for_style is None:
                    ref_content_for_style = para
                if ref_title_for_style and ref_content_for_style:
                    break
            
            # Déterminer où insérer
            if after_title:
                ref_title_para, ref_content, last_content = self._find_disclaimer_section(after_title, 1)
                if ref_title_para is None:
                    raise ValueError(f"Titre de référence introuvable : {after_title}")
                # Insérer après le dernier paragraphe de contenu de la section de référence
                insert_after = last_content if last_content else ref_title_para
                # Pour le style, privilégier la section de référence explicitement visée.
                ref_title_for_style = ref_title_para
                if ref_content:
                    ref_content_for_style = ref_content[0]
            else:
                # Insérer à la fin
                insert_after = self._get_last_disclaimer_paragraph()
                if insert_after is None:
                    # Pas de disclaimer existant, ajouter à la fin du document
                    all_paras = list(self.doc.paragraphs)
                    if all_paras:
                        insert_after = all_paras[-1]
                    else:
                        insert_after = None
            
            if insert_after is None:
                raise ValueError("Impossible de trouver où insérer la section - le document ne contient aucun paragraphe")
            
            # Créer le titre avec le style d'un titre existant
            title_para = self.doc.add_paragraph()
            title_para._element.getparent().remove(title_para._element)
            insert_after._element.addnext(title_para._element)
            
            # Copier le format de paragraphe (pPr) du titre de référence
            if ref_title_for_style:
                ref_ppr = ref_title_for_style._element.find(qn("w:pPr"))
                if ref_ppr is not None:
                    new_ppr = title_para._element.find(qn("w:pPr"))
                    if new_ppr is not None:
                        title_para._element.remove(new_ppr)
                    title_para._element.insert(0, deepcopy(ref_ppr))
            
            # Ajouter le texte du titre et copier le format de run (rPr)
            title_run = title_para.add_run(title)
            if ref_title_for_style and ref_title_for_style.runs:
                self._copy_run_format(ref_title_for_style.runs[0], title_run)
            else:
                # Fallback si pas de référence
                title_run.font.bold = True
            
            if self.markup_mode:
                title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            # Créer le contenu avec le style d'un contenu existant
            content_lines = content.split('\n')
            last_elem = title_para._element
            for line in content_lines:
                if line.strip():
                    new_p = self.doc.add_paragraph()
                    new_p._element.getparent().remove(new_p._element)
                    last_elem.addnext(new_p._element)
                    last_elem = new_p._element
                    
                    # Copier le format de paragraphe (pPr) du contenu de référence
                    if ref_content_for_style:
                        ref_ppr = ref_content_for_style._element.find(qn("w:pPr"))
                        if ref_ppr is not None:
                            new_ppr = new_p._element.find(qn("w:pPr"))
                            if new_ppr is not None:
                                new_p._element.remove(new_ppr)
                            new_p._element.insert(0, deepcopy(ref_ppr))
                    
                    # Ajouter le texte et copier le format de run (rPr)
                    run = new_p.add_run(line)
                    if ref_content_for_style and ref_content_for_style.runs:
                        self._copy_run_format(ref_content_for_style.runs[0], run)
                    
                    # Highlighting
                    if self.markup_mode:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    elif red_highlight:
                        run.font.highlight_color = WD_COLOR_INDEX.RED
        
        return self

    def remove_disclaimer_section(self, title: str, occurrence: int = 1):
        """
        Supprime une section de disclaimer complète (titre + contenu).
        En mode mark-up : barre tout au lieu de supprimer.
        """
        title_para, content_paras, _ = self._find_disclaimer_section(title, occurrence)
        if title_para is None:
            return self
        
        if self.markup_mode:
            # Barrer le titre
            for r in title_para.runs:
                r.font.strike = True
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            # Barrer le contenu
            for cp in content_paras:
                for r in cp.runs:
                    r.font.strike = True
                    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            # Supprimer le titre
            title_para._element.getparent().remove(title_para._element)
            # Supprimer le contenu
            for cp in content_paras:
                cp._element.getparent().remove(cp._element)
        
        return self

    def update_disclaimer_content(
        self,
        title: str,
        new_content: str,
        red_highlight: bool = False,
        occurrence: int = 1
    ):
        """
        Met à jour uniquement le contenu d'une section de disclaimer existante.
        Équivalent à set_disclaimer_section mais force que la section existe.
        """
        title_para, content_paras, _ = self._find_disclaimer_section(title, occurrence)
        if title_para is None:
            return self
        
        # Utiliser set_disclaimer_section qui gère déjà la mise à jour
        return self.set_disclaimer_section(title, new_content, red_highlight=red_highlight, occurrence=occurrence)

    def add_disclaimer_content(
        self,
        title: str,
        additional_content: str,
        red_highlight: bool = False,
        occurrence: int = 1
    ):
        """
        Ajoute du contenu à la fin d'une section de disclaimer existante.
        """
        title_para, content_paras, last_content = self._find_disclaimer_section(title, occurrence)
        if title_para is None:
            return self
        
        # Trouver où insérer
        insert_after = last_content if last_content else title_para
        ref_run = content_paras[0].runs[0] if content_paras and content_paras[0].runs else None
        
        # Ajouter le nouveau contenu
        content_lines = additional_content.split('\n')
        last_elem = insert_after._element
        
        for line in content_lines:
            if line.strip():
                new_p = self.doc.add_paragraph()
                new_p._element.getparent().remove(new_p._element)
                last_elem.addnext(new_p._element)
                last_elem = new_p._element
                
                # Copier le format
                if content_paras and content_paras[0]._element.find(qn("w:pPr")) is not None:
                    ppr = content_paras[0]._element.find(qn("w:pPr"))
                    new_ppr = new_p._element.find(qn("w:pPr"))
                    if new_ppr is not None:
                        new_p._element.remove(new_ppr)
                    new_p._element.insert(0, deepcopy(ppr))
                
                run = new_p.add_run(line)
                if ref_run:
                    self._copy_run_format(ref_run, run)
                
                # Highlighting
                if self.markup_mode:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                elif red_highlight:
                    run.font.highlight_color = WD_COLOR_INDEX.RED
        
        return self

    def add_logo_to_header(self, logo_path: str, width_inches: float = 1.0, all_sections: bool = True) -> bool:
        """Supprime le header existant et place uniquement le logo en haut à gauche."""
        try:
            if not self.doc.sections:
                raise ValueError("Le document ne contient aucune section")
            sections = self.doc.sections if all_sections else [self.doc.sections[0]]
            for section in sections:
                header = section.header
                # Supprimer tout le contenu du header
                for paragraph in list(header.paragraphs):
                    paragraph._element.getparent().remove(paragraph._element)
                for table in list(header.tables):
                    table._element.getparent().remove(table._element)
                # Nouveau paragraphe avec uniquement le logo
                paragraph = header.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(width_inches))
                paragraph.add_run("\n")
            return True
        except Exception as e:
            print(f"Erreur lors de l'ajout du logo : {e}")
            return False

    def save(self, output_path: str):
        self.doc.save(output_path)


# -----------------------------------------------------------------------------
# Transformer principal
# -----------------------------------------------------------------------------

class TermSheetTransformer:
    """
    Transforme un document Word term sheet et produit :
    - Document mark-up : modifications visibles (jaune + barré)
    - Document final : modifications appliquées
    - Document PDF : version PDF du final

    Fonctionnalités :
    1. Sections de tableau : Modifier les lignes du tableau principal
    2. Paragraphes non structurés : Modifier des paragraphes individuels
    3. Sections de Disclaimer : Modifier les sections titre/contenu après les tableaux
    4. Logo : Ajouter un logo dans le header

    Exemple :
        t = TermSheetTransformer("term_sheet.docx")
        
        # Remplacements de mots
        t.replace_word("Mot1", "Mot2")
        
        # Sections de tableau (méthode unifiée)
        t.set_section("Issuer", "Paul Berber", after_section="Listing")
        t.set_section("New Section", "Description")  # Ajoute à la fin si pas de référence
        t.set_section("Important", "Texte important", red_highlight_in_final=True)
        t.set_section_order(["Issuer", "Country", "Currency"])
        t.remove_section("Old Section")
        
        # Paragraphes non structurés
        t.add_content("Notice importante", red_highlight_in_final=True)
        t.update_paragraph("Ancien texte", "Nouveau texte")
        t.remove_paragraph("Paragraphe à supprimer")
        
        # Sections de Disclaimer (titre/contenu après tableaux)
        t.set_disclaimer_section("Important note", "Nouveau texte\\nLigne 2")
        t.update_disclaimer_content("Wichtiger Hinweis", "Contenu mis à jour")
        t.add_disclaimer_content("Important Risks", "Paragraphe supplémentaire")
        t.remove_disclaimer_section("Wesentliche Risiken")
        
        # Logo
        t.add_logo("logo.png")
        
        t.execute_and_export("./output", "modifications.docx", "final.docx", "final.pdf")
    """

    def __init__(self, docx_path: str):
        self._temp_dir: Optional[Path] = None
        path = Path(docx_path)
        if not path.exists():
            raise FileNotFoundError(f"Fichier introuvable : {docx_path}")
        if path.suffix.lower() == ".doc":
            self.docx_path, self._temp_dir = _convert_doc_to_docx(path)
        else:
            self.docx_path = path
        self.operations: List = []

    def __del__(self):
        if hasattr(self, "_temp_dir") and self._temp_dir is not None and self._temp_dir.exists():
            try:
                shutil.rmtree(self._temp_dir, ignore_errors=True)
            except Exception:
                pass

    def get_section_description(self, section_title: str, occurrence: int = 1) -> Optional[str]:
        """
        Lit la description d'une section directement dans le document source et la retourne.

        Retourne None si la section est introuvable.

        Pour les tableaux à deux colonnes sans fusion verticale : les lignes suivantes dont la
        colonne titre est vide sont concaténées (textes des cellules B reliés par des sauts de ligne).

        Exemple :
            issuer = transformer.get_section_description("Issuer")
            transformer.set_section("Issuer", issuer + " bank")
        """
        doc = Document(str(self.docx_path))
        editor = _TermSheetEditor(doc, markup_mode=False)
        return editor.get_section_description(section_title, occurrence)

    def replace_word(self, old: str, new: str, occurrence: int = 1) -> "TermSheetTransformer":
        self.operations.append(ReplaceOp(old=old, new=new, occurrence=occurrence))
        return self

    def add_section_after(self, after_section: str, title: str, description: str,
                         occurrence: int = 1) -> "TermSheetTransformer":
        self.operations.append(AddSectionOp(
            after_section=after_section, title=title, description=description, occurrence=occurrence
        ))
        return self

    def update_section_description(self, section_title: str, new_description: str,
                                   occurrence: int = 1) -> "TermSheetTransformer":
        self.operations.append(UpdateDescriptionOp(
            section_title=section_title, new_description=new_description, occurrence=occurrence
        ))
        return self

    def set_section(
        self, 
        title: str, 
        description: str,
        after_section: Optional[str] = None,
        red_highlight_in_final: bool = False,
        occurrence: int = 1
    ) -> "TermSheetTransformer":
        """
        Ajoute ou met à jour une section dans le tableau.
        - Si la section existe : met à jour la description
        - Si la section n'existe pas : l'ajoute après after_section (ou à la fin si None)
        - red_highlight_in_final : surligne en rouge dans la version finale (pas dans le markup)
        """
        self.operations.append(SetSectionOp(
            title=title,
            description=description,
            after_section=after_section,
            red_highlight_in_final=red_highlight_in_final,
            occurrence=occurrence
        ))
        return self

    def set_section_order(self, section_order: List[str]) -> "TermSheetTransformer":
        """
        Définit l'ordre des sections dans le tableau.
        Les sections non listées sont placées à la fin.
        """
        self.operations.append(SetSectionOrderOp(section_order=section_order))
        return self

    def remove_section(self, section_title: str, occurrence: int = 1) -> "TermSheetTransformer":
        self.operations.append(DeleteSectionOp(section_title=section_title, occurrence=occurrence))
        return self

    def add_logo(self, logo_path: str, width_inches: float = 1.0,
                 all_sections: bool = True) -> "TermSheetTransformer":
        self.operations.append(AddLogoOp(
            logo_path=logo_path, width_inches=width_inches, all_sections=all_sections
        ))
        return self

    def add_content(self, text: str, after_paragraph: Optional[str] = None,
                    red_highlight_in_final: bool = False) -> "TermSheetTransformer":
        """Ajoute un titre ou une description (paragraphe hors tableau)."""
        self.operations.append(AddContentOp(
            text=text, after_paragraph=after_paragraph, red_highlight_in_final=red_highlight_in_final
        ))
        return self

    def remove_paragraph(self, text_contains: str, occurrence: int = 1) -> "TermSheetTransformer":
        """Supprime un paragraphe contenant le texte donné."""
        self.operations.append(RemoveParagraphOp(text_contains=text_contains, occurrence=occurrence))
        return self

    def update_paragraph(self, text_contains: str, new_text: str,
                         occurrence: int = 1) -> "TermSheetTransformer":
        """Modifie le contenu d'un paragraphe."""
        self.operations.append(UpdateParagraphOp(
            text_contains=text_contains, new_text=new_text, occurrence=occurrence
        ))
        return self

    # -------------------------------------------------------------------------
    # Gestion des sections de Disclaimer (paragraphes structurés après les tableaux)
    # -------------------------------------------------------------------------

    def set_disclaimer_section(
        self,
        title: str,
        content: str,
        after_title: Optional[str] = None,
        red_highlight_in_final: bool = False,
        occurrence: int = 1
    ) -> "TermSheetTransformer":
        """
        Ajoute ou met à jour une section dans la partie Disclaimer (après les tableaux).
        
        - Si la section existe : remplace son contenu
        - Si la section n'existe pas : l'ajoute après after_title (ou à la fin si None)
        - Le content peut contenir des \\n pour séparer les paragraphes
        - red_highlight_in_final : surligne en rouge dans la version finale
        
        Exemple:
            t.set_disclaimer_section("Important note", "Texte ligne 1\\nTexte ligne 2")
            t.set_disclaimer_section("New Warning", "Attention!", after_title="Important note")
        """
        self.operations.append(SetDisclaimerSectionOp(
            title=title,
            content=content,
            after_title=after_title,
            red_highlight_in_final=red_highlight_in_final,
            occurrence=occurrence
        ))
        return self

    def remove_disclaimer_section(self, title: str, occurrence: int = 1) -> "TermSheetTransformer":
        """
        Supprime une section de disclaimer complète (titre + contenu).
        
        Exemple:
            t.remove_disclaimer_section("Important note")
        """
        self.operations.append(RemoveDisclaimerSectionOp(title=title, occurrence=occurrence))
        return self

    def update_disclaimer_content(
        self,
        title: str,
        new_content: str,
        red_highlight_in_final: bool = False,
        occurrence: int = 1
    ) -> "TermSheetTransformer":
        """
        Met à jour uniquement le contenu d'une section de disclaimer existante.
        La section doit exister, sinon une erreur sera levée.
        
        Équivalent à set_disclaimer_section mais force que la section existe.
        
        Exemple:
            t.update_disclaimer_content("Important note", "Nouveau contenu\\nDeuxième ligne")
        """
        self.operations.append(UpdateDisclaimerContentOp(
            title=title,
            new_content=new_content,
            red_highlight_in_final=red_highlight_in_final,
            occurrence=occurrence
        ))
        return self

    def add_disclaimer_content(
        self,
        title: str,
        additional_content: str,
        red_highlight_in_final: bool = False,
        occurrence: int = 1
    ) -> "TermSheetTransformer":
        """
        Ajoute du contenu à la fin d'une section de disclaimer existante.
        La section doit exister, sinon une erreur sera levée.
        
        Exemple:
            t.add_disclaimer_content("Important note", "Paragraphe supplémentaire")
        """
        self.operations.append(AddDisclaimerContentOp(
            title=title,
            additional_content=additional_content,
            red_highlight_in_final=red_highlight_in_final,
            occurrence=occurrence
        ))
        return self

    def execute_and_export(
        self,
        output_dir: str,
        markup_docx: str = "modifications.docx",
        final_docx: str = "document_final.docx",
        final_pdf: str = "document_final.pdf",
    ) -> None:
        """
        Exécute toutes les modifications et exporte les 3 fichiers.
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        doc_markup = Document(str(self.docx_path))
        doc_final = Document(str(self.docx_path))

        editor_markup = _TermSheetEditor(doc_markup, markup_mode=True)
        editor_final = _TermSheetEditor(doc_final, markup_mode=False)

        for op in self.operations:
            if isinstance(op, ReplaceOp):
                editor_markup.replace_text(op.old, op.new)
                editor_final.replace_text(op.old, op.new)
            elif isinstance(op, AddSectionOp):
                editor_markup.insert_section_after(
                    op.after_section, op.title, op.description, op.occurrence
                )
                editor_final.insert_section_after(
                    op.after_section, op.title, op.description, op.occurrence
                )
            elif isinstance(op, UpdateDescriptionOp):
                editor_markup.update_section_description(
                    op.section_title, op.new_description, op.occurrence
                )
                editor_final.update_section_description(
                    op.section_title, op.new_description, op.occurrence
                )
            elif isinstance(op, SetSectionOp):
                editor_markup.set_section(
                    op.title, op.description, op.after_section, 
                    red_highlight=False, occurrence=op.occurrence
                )
                editor_final.set_section(
                    op.title, op.description, op.after_section,
                    red_highlight=op.red_highlight_in_final, occurrence=op.occurrence
                )
            elif isinstance(op, SetSectionOrderOp):
                editor_markup.set_section_order(op.section_order)
                editor_final.set_section_order(op.section_order)
            elif isinstance(op, DeleteSectionOp):
                editor_markup.delete_section(op.section_title, op.occurrence)
                editor_final.delete_section(op.section_title, op.occurrence)
            elif isinstance(op, AddLogoOp):
                editor_markup.add_logo_to_header(
                    op.logo_path, op.width_inches, op.all_sections
                )
                editor_final.add_logo_to_header(
                    op.logo_path, op.width_inches, op.all_sections
                )
            elif isinstance(op, AddContentOp):
                editor_markup.add_content(
                    op.text, op.after_paragraph,
                    highlight=True,  # jaune en mark-up
                    red_highlight=False,
                )
                editor_final.add_content(
                    op.text, op.after_paragraph,
                    highlight=False,
                    red_highlight=op.red_highlight_in_final,
                )
            elif isinstance(op, RemoveParagraphOp):
                editor_markup.remove_paragraph(op.text_contains, op.occurrence)
                editor_final.remove_paragraph(op.text_contains, op.occurrence)
            elif isinstance(op, UpdateParagraphOp):
                editor_markup.update_paragraph(op.text_contains, op.new_text, op.occurrence)
                editor_final.update_paragraph(op.text_contains, op.new_text, op.occurrence)
            elif isinstance(op, SetDisclaimerSectionOp):
                editor_markup.set_disclaimer_section(
                    op.title, op.content, op.after_title,
                    red_highlight=False, occurrence=op.occurrence
                )
                editor_final.set_disclaimer_section(
                    op.title, op.content, op.after_title,
                    red_highlight=op.red_highlight_in_final, occurrence=op.occurrence
                )
            elif isinstance(op, RemoveDisclaimerSectionOp):
                editor_markup.remove_disclaimer_section(op.title, op.occurrence)
                editor_final.remove_disclaimer_section(op.title, op.occurrence)
            elif isinstance(op, UpdateDisclaimerContentOp):
                editor_markup.update_disclaimer_content(
                    op.title, op.new_content,
                    red_highlight=False, occurrence=op.occurrence
                )
                editor_final.update_disclaimer_content(
                    op.title, op.new_content,
                    red_highlight=op.red_highlight_in_final, occurrence=op.occurrence
                )
            elif isinstance(op, AddDisclaimerContentOp):
                editor_markup.add_disclaimer_content(
                    op.title, op.additional_content,
                    red_highlight=False, occurrence=op.occurrence
                )
                editor_final.add_disclaimer_content(
                    op.title, op.additional_content,
                    red_highlight=op.red_highlight_in_final, occurrence=op.occurrence
                )

        path_markup = output_dir / markup_docx
        path_final = output_dir / final_docx
        path_pdf = output_dir / final_pdf

        editor_markup.save(str(path_markup))
        _strip_all_highlights(doc_final)
        editor_final.save(str(path_final))

        try:
            import docx2pdf
            docx2pdf.convert(str(path_final), str(path_pdf))
        except Exception as e:
            print(f"Erreur lors de la conversion PDF : {e}")
            print("Assurez-vous que Word est installé et que docx2pdf est correctement configuré.")

        print(f"Mark-up sauvegardé : {path_markup}")
        print(f"Document final sauvegardé : {path_final}")
        print(f"PDF sauvegardé : {path_pdf}")


# -----------------------------------------------------------------------------
# Exemple d'utilisation
# -----------------------------------------------------------------------------

if __name__ == "__main__":
    # Accepte .doc (conversion auto via LibreOffice) ou .docx
    transformer = TermSheetTransformer("mon_term_sheet.docx")

    # Remplacements de mots
    transformer.replace_word("Mot1", "Mot2")
    
    # Nouvelle méthode set_section : ajoute ou met à jour une section
    # Si "Issuer" n'existe pas, l'ajoute après "Listing"
    # Si "Issuer" existe, met à jour sa description
    transformer.set_section("Issuer", "Paul Berber", after_section="Listing")
    
    # Si after_section n'est pas spécifié, ajoute à la fin du tableau
    transformer.set_section("New Section", "Description de la nouvelle section")
    
    # Avec surlignage rouge dans la version finale
    transformer.set_section("Important Note", "Texte important", red_highlight_in_final=True)
    
    # Pour des retours à la ligne dans la description, utiliser \n
    transformer.set_section("Multi-line", "Ligne 1\nLigne 2\nLigne 3")
    
    # Anciennes méthodes toujours disponibles
    transformer.add_section_after("Listing", "Other Section", "Other Description")
    transformer.update_section_description("Country", "France")
    
    # Définir l'ordre des sections
    transformer.set_section_order(["Issuer", "Country", "Currency", "Amount"])
    
    # Suppression
    transformer.remove_section("Old Section")
    
    # Contenu hors tableau (paragraphes non structurés)
    transformer.add_content("Notice importante", red_highlight_in_final=True)
    transformer.add_content("Paragraphe après un autre", after_paragraph="Texte existant")
    transformer.update_paragraph("Ancien texte", "Nouveau texte")
    transformer.remove_paragraph("Paragraphe à supprimer")
    
    # -------------------------------------------------------------------------
    # Sections de Disclaimer (paragraphes structurés titre/contenu après tableaux)
    # -------------------------------------------------------------------------
    
    # Modifier une section de disclaimer existante
    transformer.set_disclaimer_section(
        "Important note",
        "Nouveau texte pour cette section\nDeuxième paragraphe"
    )
    
    # Ajouter une nouvelle section de disclaimer
    transformer.set_disclaimer_section(
        "New Warning",
        "Ceci est un nouveau disclaimer\nAvec plusieurs lignes",
        after_title="Important note",
        red_highlight_in_final=True
    )
    
    # Mettre à jour uniquement le contenu (la section doit exister)
    transformer.update_disclaimer_content(
        "Wichtiger Hinweis",
        "Contenu mis à jour"
    )
    
    # Ajouter du contenu à la fin d'une section existante
    transformer.add_disclaimer_content(
        "Important Risks",
        "Paragraphe supplémentaire ajouté"
    )
    
    # Supprimer une section de disclaimer complète
    transformer.remove_disclaimer_section("Wesentliche Risiken")
    
    # Logo
    transformer.add_logo("logo_nbc.png", width_inches=0.8)

    transformer.execute_and_export(
        output_dir="./output",
        markup_docx="modifications.docx",
        final_docx="document_final.docx",
        final_pdf="document_final.pdf",
    )
